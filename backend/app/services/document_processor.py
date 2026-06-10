import PyPDF2
from docx import Document
import os


class DocumentProcessor:
    """Handle document extraction for PDF and DOCX files"""
    
    @staticmethod
    def extract_text(file_path: str) -> str:
        """
        Extract text from PDF or DOCX file
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == ".pdf":
            return DocumentProcessor._extract_from_pdf(file_path)
        elif file_extension == ".docx":
            return DocumentProcessor._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            with open(file_path, "rb") as pdf_file:
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text()
                    text += "\n"
        
        except Exception as e:
            raise Exception(f"Error extracting PDF: {str(e)}")
        
        return text.strip()
    
    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(file_path)
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
        
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")
        
        return text.strip()
